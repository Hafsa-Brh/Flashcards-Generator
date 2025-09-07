"""
Document loader for various file formats.

This module handles loading and extracting text from different file types:
- PDF files (.pdf)
- Word documents (.docx)
- Text files (.txt)
- Markdown files (.md)
- HTML files (.html)
"""

import logging
import mimetypes
from pathlib import Path
from typing import Optional, Dict, Any, List
from uuid import uuid4

import pypdf
from docx import Document
import markdown

from ..schemas import Source, SourceType
from ..config import get_settings


logger = logging.getLogger(__name__)


class DocumentLoader:
    """Handles loading documents from various file formats."""
    
    def __init__(self):
        self.settings = get_settings()
        
        # MIME type to SourceType mapping
        self.mime_mapping = {
            'application/pdf': SourceType.PDF,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': SourceType.DOCX,
            'text/plain': SourceType.TXT,
            'text/markdown': SourceType.MARKDOWN,
            'text/html': SourceType.HTML,
        }
        
        # File extension to SourceType mapping (fallback)
        self.extension_mapping = {
            '.pdf': SourceType.PDF,
            '.docx': SourceType.DOCX,
            '.doc': SourceType.DOCX,  # Treat legacy .doc as .docx
            '.txt': SourceType.TXT,
            '.md': SourceType.MARKDOWN,
            '.markdown': SourceType.MARKDOWN,
            '.html': SourceType.HTML,
            '.htm': SourceType.HTML,
        }
    
    def detect_file_type(self, file_path: Path) -> SourceType:
        """Detect the file type based on MIME type and extension."""
        
        # Try MIME type detection first
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if mime_type in self.mime_mapping:
            return self.mime_mapping[mime_type]
        
        # Fallback to extension-based detection
        extension = file_path.suffix.lower()
        if extension in self.extension_mapping:
            return self.extension_mapping[extension]
        
        # Default to text if unknown
        logger.warning(f"Unknown file type for {file_path}, treating as text")
        return SourceType.TXT
    
    def load_pdf(self, file_path: Path) -> str:
        """Load text content from a PDF file."""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                if len(pdf_reader.pages) == 0:
                    raise ValueError("PDF file has no pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                        else:
                            logger.warning(f"No text extracted from page {page_num + 1}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
            
            if not text_content:
                raise ValueError("No text could be extracted from PDF")
            
            full_text = "\n\n".join(text_content)
            logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise
    
    def load_docx(self, file_path: Path) -> str:
        """Load text content from a DOCX file."""
        try:
            doc = Document(file_path)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_content.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                raise ValueError("No text content found in DOCX file")
            
            full_text = "\n\n".join(text_content)
            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise
    
    def load_text(self, file_path: Path) -> str:
        """Load content from a plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        logger.info(f"Successfully loaded text file with {encoding} encoding")
                        return content
                except UnicodeDecodeError:
                    continue
            
            raise ValueError(f"Could not decode text file with any of: {encodings}")
            
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            raise
    
    def load_markdown(self, file_path: Path) -> str:
        """Load content from a Markdown file."""
        try:
            # Load the raw markdown content first
            raw_content = self.load_text(file_path)
            
            # Convert markdown to plain text (remove formatting)
            html = markdown.markdown(raw_content)
            
            # Basic HTML tag removal (you could use BeautifulSoup for better parsing)
            import re
            text_content = re.sub(r'<[^>]+>', '', html)
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            logger.info(f"Successfully processed markdown file")
            return text_content
            
        except Exception as e:
            logger.error(f"Error loading markdown file {file_path}: {e}")
            raise
    
    def load_html(self, file_path: Path) -> str:
        """Load content from an HTML file."""
        try:
            from bs4 import BeautifulSoup
            
            html_content = self.load_text(file_path)
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text_content = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text_content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text_content = ' '.join(chunk for chunk in chunks if chunk)
            
            logger.info(f"Successfully processed HTML file")
            return text_content
            
        except ImportError:
            logger.warning("BeautifulSoup not available, using basic HTML processing")
            # Fallback: basic HTML tag removal
            html_content = self.load_text(file_path)
            import re
            text_content = re.sub(r'<[^>]+>', '', html_content)
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            return text_content
            
        except Exception as e:
            logger.error(f"Error loading HTML file {file_path}: {e}")
            raise
    
    def load_document(self, file_path: Path, title: Optional[str] = None) -> Source:
        """Load a document and return a Source object."""
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")
        
        # Detect file type
        source_type = self.detect_file_type(file_path)
        
        # Generate title if not provided
        if title is None:
            title = file_path.stem
        
        logger.info(f"Loading {source_type.value} file: {file_path}")
        
        # Load content based on file type
        try:
            if source_type == SourceType.PDF:
                content = self.load_pdf(file_path)
            elif source_type == SourceType.DOCX:
                content = self.load_docx(file_path)
            elif source_type == SourceType.TXT:
                content = self.load_text(file_path)
            elif source_type == SourceType.MARKDOWN:
                content = self.load_markdown(file_path)
            elif source_type == SourceType.HTML:
                content = self.load_html(file_path)
            else:
                raise ValueError(f"Unsupported source type: {source_type}")
            
            # Create Source object
            source = Source(
                title=title,
                file_path=str(file_path),
                source_type=source_type,
                content=content,
                metadata={
                    "file_size": file_path.stat().st_size,
                    "character_count": len(content),
                    "word_count": len(content.split()),
                }
            )
            
            logger.info(f"Successfully loaded document: {title} ({len(content)} chars)")
            return source
            
        except Exception as e:
            logger.error(f"Failed to load document {file_path}: {e}")
            raise
    
    def load_documents(self, file_paths: List[Path]) -> List[Source]:
        """Load multiple documents."""
        sources = []
        errors = []
        
        for file_path in file_paths:
            try:
                source = self.load_document(file_path)
                sources.append(source)
            except Exception as e:
                errors.append(f"{file_path}: {str(e)}")
                logger.error(f"Failed to load {file_path}: {e}")
        
        if errors:
            logger.warning(f"Failed to load {len(errors)} documents: {errors}")
        
        logger.info(f"Successfully loaded {len(sources)} out of {len(file_paths)} documents")
        return sources
    
    def load_from_directory(
        self, 
        directory: Path, 
        recursive: bool = True,
        file_patterns: Optional[List[str]] = None
    ) -> List[Source]:
        """Load all supported documents from a directory."""
        
        if file_patterns is None:
            file_patterns = ['*.pdf', '*.docx', '*.doc', '*.txt', '*.md', '*.html']
        
        file_paths = []
        
        for pattern in file_patterns:
            if recursive:
                file_paths.extend(directory.rglob(pattern))
            else:
                file_paths.extend(directory.glob(pattern))
        
        logger.info(f"Found {len(file_paths)} files in {directory}")
        
        return self.load_documents(file_paths)


# Convenience functions
def load_document(file_path: Path, title: Optional[str] = None) -> Source:
    """Load a single document."""
    loader = DocumentLoader()
    return loader.load_document(file_path, title)


def load_documents(file_paths: List[Path]) -> List[Source]:
    """Load multiple documents."""
    loader = DocumentLoader()
    return loader.load_documents(file_paths)


def load_from_directory(
    directory: Path, 
    recursive: bool = True,
    file_patterns: Optional[List[str]] = None
) -> List[Source]:
    """Load all documents from a directory."""
    loader = DocumentLoader()
    return loader.load_from_directory(directory, recursive, file_patterns)
